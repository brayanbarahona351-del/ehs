import streamlit as st
import pandas as pd
from docx import Document
from io import BytesIO

# --- CONFIGURACIÓN Y DATOS TÉCNICOS ---
st.set_page_config(page_title="Escala Goldstein - Psicología", layout="wide")

# Grupos según el Manual de Ambrosio Tomás
GRUPOS = {
    "I: Primeras Habilidades": list(range(1, 9)),
    "II: Habilidades Avanzadas": list(range(9, 15)),
    "III: Habilidades Relacionadas con los Sentimientos": list(range(15, 22)),
    "IV: Habilidades Alternativas a la Agresión": list(range(22, 31)),
    "V: Habilidades para hacer frente al Estrés": list(range(31, 43)),
    "VI: Habilidades de Planificación": list(range(43, 51))
}

# Baremos Eneatípicos (Muestra Universitaria de Psicología - Manual pág. 7)
BAREMOS = {
    "I":   {9: 38, 8: 35, 7: 33, 6: 30, 5: 28, 4: 26, 3: 23, 2: 21, 1: 0},
    "II":  {9: 28, 8: 26, 7: 24, 6: 22, 5: 21, 4: 19, 3: 17, 2: 15, 1: 0},
    "III": {9: 33, 8: 31, 7: 29, 6: 26, 5: 24, 4: 22, 3: 20, 2: 17, 1: 0},
    "IV":  {9: 43, 8: 41, 7: 38, 6: 36, 5: 33, 4: 31, 3: 29, 2: 26, 1: 0},
    "V":   {9: 56, 8: 53, 7: 49, 6: 46, 5: 42, 4: 39, 3: 35, 2: 32, 1: 0},
    "VI":  {9: 40, 8: 38, 7: 35, 6: 33, 5: 31, 4: 28, 3: 26, 2: 23, 1: 0},
    "TOTAL": {9: 228, 8: 216, 7: 204, 6: 192, 5: 181, 4: 169, 3: 157, 2: 145, 1: 0}
}

PREGUNTAS = [
    "Prestas atención a la persona que te está hablando...", "Hablas con los demás de temas poco importantes...",
    "Hablas con otras personas sobre cosas que interesan a ambos", "Clarificas la información que necesitas...",
    "Permites que los demás sepan que les agradeces los favores", "Te das a conocer a los demás por propia iniciativa",
    "Ayudas a que los demás se conozcan entre sí", "Dices que te gusta algún aspecto de la otra persona...",
    "Pides que te ayuden cuando tienes alguna dificultad", "Eliges la mejor forma para integrarte en un grupo...",
    "Explicas con claridad a los demás cómo hacer una tarea...", "Prestas atención a las instrucciones...",
    "Pides disculpas a los demás por haber hecho algo mal", "Intentas persuadir a los demás de que tus ideas son mejores...",
    "Intentas reconocer las emociones que experimentas", "Permites que los demás conozcan lo que sientes",
    "Intentas comprender lo que sienten los demás", "Intentas comprender el enfado de la otra persona",
    "Permites que los demás sepan que te interesas por ellos", "Piensas porqué estás asustado y haces algo para disminuirlo",
    "Te dices cosas agradables cuando mereces recompensa", "Reconoces cuando es necesario pedir permiso...",
    "Te ofreces para compartir algo apreciado por los demás", "Ayudas a quien lo necesita",
    "Estableces un sistema de negociación satisfactorio", "Controlas tu carácter para no perder el control",
    "Defiendes tus derechos dando a conocer tu postura", "Te las arreglas sin perder el control ante bromas",
    "Te mantienes al margen de situaciones problemáticas", "Encuentras formas de resolver conflictos sin pelear",
    "Dices a los demás cuándo han sido responsables de un problema", "Intentas llegar a una solución justa ante una queja",
    "Expresas un sincero cumplido por cómo han jugado", "Haces algo para sentir menos vergüenza",
    "Eres consciente si te dejan de lado y buscas sentirte mejor", "Manifiestas que han tratado injustamente a un amigo",
    "Consideras la posición de la otra persona antes de decidir", "Comprendes por qué has fracasado y cómo mejorar",
    "Resuelves la confusión ante mensajes contradictorios", "Comprendes una acusación y piensas cómo relacionarte",
    "Planificas la mejor forma para exponer tu punto de vista", "Decides lo que quieres hacer frente a la presión grupal",
    "Resuelves la sensación de aburrimiento...", "Reconoces si un evento está bajo tu control",
    "Tomas decisiones realistas sobre tus capacidades", "Eres realista sobre cómo desenvolverte en una tarea",
    "Resuelves qué necesitas saber y cómo conseguir la info", "Determinas cuál problema es el más importante",
    "Consideras posibilidades y eliges la mejor", "Te organizas y preparas para facilitar tu trabajo"
]

# --- FUNCIONES DE APOYO ---
def obtener_eneatipo(puntaje, clave):
    tabla = BAREMOS[clave]
    for enea in sorted(tabla.keys(), reverse=True):
        if puntaje >= tabla[enea]: return enea
    return 1

def interpretar(enea):
    if enea >= 7: return "Nivel Competente (Alto)"
    if enea >= 4: return "Nivel Promedio"
    return "Nivel Deficitario (Bajo)"

def crear_word(nombre, edad, resultados, total_pd, enea_t, diag_g, respuestas_list):
    doc = Document()
    doc.add_heading('INFORME PSICOLÓGICO: HABILIDADES SOCIALES', 0)
    
    table_info = doc.add_table(rows=2, cols=2)
    table_info.cell(0,0).text = f"Nombre: {nombre}"
    table_info.cell(0,1).text = f"Edad: {edad}"
    table_info.cell(1,0).text = "Instrumento: Lista de Goldstein"
    table_info.cell(1,1).text = "Baremo: Estudiantes de Psicología"

    doc.add_heading('Resultados por Áreas', level=1)
    t = doc.add_table(rows=1, cols=4)
    t.style = 'Table Grid'
    cols = ['Área', 'PD', 'Eneatipo', 'Diagnóstico']
    for i, texto in enumerate(cols): t.rows[0].cells[i].text = texto

    for res in resultados:
        row = t.add_row().cells
        row[0].text = res['Área']
        row[1].text = str(res['Puntaje'])
        row[2].text = str(res['Eneatipo'])
        row[3].text = res['Diagnóstico']

    doc.add_heading('Interpretación General', level=1)
    p = doc.add_paragraph()
    p.add_run(f"Puntaje Directo Total: {total_pd}\nEneatipo Global: {enea_t}\nDiagnóstico General: {diag_g}").bold = True

    # Guardar en buffer
    target = BytesIO()
    doc.save(target)
    return target.getvalue()

# --- INTERFAZ ---
st.title("Sistema de Evaluación de Habilidades Sociales")
with st.sidebar:
    st.header("Datos del Evaluado")
    nombre = st.text_input("Nombre completo")
    edad = st.number_input("Edad", 12, 90, 20)

st.subheader("Por favor, responda a cada ítem según la frecuencia de su conducta")
respuestas = {}
for i, texto in enumerate(PREGUNTAS, 1):
    respuestas[i] = st.radio(f"{i}. {texto}", [1,2,3,4,5], format_func=lambda x: {1:"Nunca", 2:"Pocas veces", 3:"A veces", 4:"Muchas veces", 5:"Siempre"}[x], horizontal=True, key=f"p{i}")

if st.button("Procesar e Interpretar"):
    res_finales = []
    total_pd = sum(respuestas.values())
    
    for g_nombre, items in GRUPOS.items():
        pd_grupo = sum(respuestas[idx] for idx in items)
        clave = g_nombre.split(":")[0].split(" ")[1]
        enea = obtener_eneatipo(pd_grupo, clave)
        res_final
