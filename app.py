import streamlit as st
import pandas as pd
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO
import matplotlib.pyplot as plt

# --- CONFIGURACIÓN DE PÁGINA ---
st.set_page_config(page_title="EHS Goldstein - Clínica Avanzada", layout="wide")

if 'revision' not in st.session_state:
    st.session_state.revision = 0

def reset_form():
    st.session_state.revision += 1
    st.rerun()

# --- DATOS TÉCNICOS ---
GRUPOS = {
    "I: Primeras HHSS": {"items": list(range(1, 9)), "desc": "Habilidades de iniciación social y cordialidad básica."},
    "II: HHSS Avanzadas": {"items": list(range(9, 15)), "desc": "Asertividad para integrarse y seguir o dar directrices."},
    "III: HHSS de Sentimientos": {"items": list(range(15, 22)), "desc": "Inteligencia emocional, empatía y auto-reforzamiento."},
    "IV: Alt. a la Agresión": {"items": list(range(22, 31)), "desc": "Control de impulsos, negociación y resolución de problemas."},
    "V: Frente al Estrés": {"items": list(range(31, 43)), "desc": "Manejo de la frustración, vergüenza y presión social."},
    "VI: Planificación": {"items": list(range(43, 51)), "desc": "Funciones ejecutivas: toma de decisiones y organización."}
}

BAREMOS = {
    "I":   {9: 38, 8: 35, 7: 33, 6: 30, 5: 28, 4: 26, 3: 23, 2: 21, 1: 0},
    "II":  {9: 28, 8: 26, 7: 24, 6: 22, 5: 21, 4: 19, 3: 17, 2: 15, 1: 0},
    "III": {9: 33, 8: 31, 7: 29, 6: 26, 5: 24, 4: 22, 3: 20, 2: 17, 1: 0},
    "IV":  {9: 43, 8: 41, 7: 38, 6: 36, 5: 33, 4: 31, 3: 29, 2: 26, 1: 0},
    "V":   {9: 56, 8: 53, 7: 49, 6: 46, 5: 42, 4: 39, 3: 35, 2: 32, 1: 0},
    "VI":  {9: 40, 8: 38, 7: 35, 6: 33, 5: 31, 4: 28, 3: 26, 2: 23, 1: 0},
    "TOTAL": {9: 228, 8: 216, 7: 204, 6: 192, 5: 181, 4: 169, 3: 157, 2: 145, 1: 0}
}

PREGUNTAS = [
    "Prestas atención a quien te habla", "Hablas de temas poco importantes", "Hablas de cosas que interesan a ambos",
    "Clarificas la información que necesitas", "Agradeces los favores", "Te das a conocer por propia iniciativa",
    "Ayudas a que los demás se conozcan", "Dices que te gusta algo de otra persona", "Pides ayuda ante dificultades",
    "Eliges la mejor forma de integrarte", "Explicas con claridad cómo hacer tareas", "Prestas atención a instrucciones",
    "Pides disculpas por errores", "Intentas persuadir a los demás", "Reconoces tus emociones",
    "Permites que otros sepan lo que sientes", "Intentas comprender lo que sienten otros", "Comprendes el enfado ajeno",
    "Te interesas por los demás", "Buscas disminuir tus miedos", "Te dices cosas agradables",
    "Pides permiso cuando es necesario", "Compartes algo apreciado", "Ayudas a quien lo necesita",
    "Negocias de forma satisfactoria", "Controlas tu carácter", "Defiendes tus derechos", "No pierdes el control ante bromas",
    "Evitas situaciones problemáticas", "Resuelves conflictos sin pelear", "Indicas quién es responsable del problema",
    "Buscas soluciones justas ante quejas", "Expresas cumplidos sinceros", "Haces algo para sentir menos vergüenza",
    "Gestionas cuando te dejan de lado", "Defiendes a amigos ante injusticias", "Consideras la posición de la otra persona",
    "Comprendes por qué has fracasado", "Resuelves mensajes contradictorios", "Comprendes acusaciones recibidas",
    "Planificas cómo exponer tu punto de vista", "Decides qué hacer ante presión grupal", "Resuelves el aburrimiento",
    "Reconoces si un evento está bajo tu control", "Eres realista sobre tus capacidades", "Eres realista ante tareas difíciles",
    "Resuelves qué necesitas saber", "Determinas qué problema es prioritario", "Consideras posibilidades y eliges", "Te organizas para facilitar el trabajo"
]

# --- LÓGICA DE DIAGNÓSTICO EXTENSO Y EMOJIS ---
def obtener_diagnostico(enea):
    if enea >= 7:
        return {
            "nivel": "Excelente Nivel", "emoji": "😃", "simbolo": "🌟",
            "analisis": "El evaluado muestra un repertorio conductual altamente adaptativo. Posee fluidez, asertividad y capacidad para modelar estas conductas. Es capaz de gestionar sus emociones y relacionarse de forma prosocial sin experimentar ansiedad limitante."
        }
    elif enea >= 4:
        return {
            "nivel": "Nivel Promedio", "emoji": "🙂", "simbolo": "⚖️",
            "analisis": "El evaluado posee una competencia social dentro de la norma. Logra emitir las conductas esperadas en situaciones de baja o media tensión, pero podría recurrir a la evitación o experimentar bloqueos ante estresores psicosociales elevados o conflictos prolongados."
        }
    else:
        return {
            "nivel": "Nivel Deficitario", "emoji": "😟", "simbolo": "🚨",
            "analisis": "Existe un déficit clínico significativo. Las causas probables incluyen: altos niveles de ansiedad social (fobia social), falta de modelos de aprendizaje adecuados durante el desarrollo, o un historial de contingencias punitivas al intentar socializar. Se requiere intervención prioritaria y reestructuración cognitiva."
        }

# --- INTERFAZ: BARRA LATERAL FIJA ---
with st.sidebar:
    st.header("📋 Datos e Instrucciones")
    nombre = st.text_input("Nombre Completo", key=f"n_{st.session_state.revision}")
    edad = st.number_input("Edad", 12, 99, 20, key=f"e_{st.session_state.revision}")
    consent = st.checkbox("✅ Acepto el Consentimiento Informado")
    
    st.divider()
    st.warning("""
    **GUÍA DE LLENADO:**
    Lea cada ítem y seleccione cómo actúa usted habitualmente.
    * **1:** Nunca / Rara vez
    * **2:** Pocas veces
    * **3:** A veces
    * **4:** Muchas veces
    * **5:** Siempre / Casi siempre
    """)
    if st.button("🗑️ Reiniciar Evaluación"): reset_form()

# --- INTERFAZ: ÁREA PRINCIPAL ---
st.title("Sistema Clínico de Evaluación: Escala de Goldstein")

if consent:
    st.write("Complete el cuestionario a continuación desplazándose hacia abajo:")
    
    respuestas = {}
    for i, p in enumerate(PREGUNTAS, 1):
        respuestas[i] = st.radio(f"**{i}. {p}**", [1,2,3,4,5], horizontal=True, key=f"q_{i}_{st.session_state.revision}", index=None)

    if st.button("📈 GENERAR ANÁLISIS Y PLAN TERAPÉUTICO"):
        if None in respuestas.values():
            st.error("⚠️ Faltan preguntas por responder. Revise el formulario.")
        else:
            # CÁLCULOS
            res_data = []
            total_pd = sum(respuestas.values())
            
            def calcular_enea(pd, k):
                for e, lim in sorted(BAREMOS[k].items(), reverse=True):
                    if pd >= lim: return e
                return 1

            for g, info in GRUPOS.items():
                pd_g = sum(respuestas[idx] for idx in info["items"])
                enea = calcular_enea(pd_g, g.split(":")[0])
                diag = obtener_diagnostico(enea)
                res_data.append({
                    "Área": g, "PD": pd_g, "Eneatipo": enea, 
                    "Emoji": diag['emoji'], "Estado": diag['nivel'], 
                    "Analisis": diag['analisis'], "Desc": info['desc']
                })

            df = pd.DataFrame(res_data)
            enea_total = calcular_enea(total_pd, "TOTAL")
            diag_total = obtener_diagnostico(enea_total)

            # --- RESULTADOS EN PANTALLA ---
            st.divider()
            st.header(f"Resultados de {nombre} {diag_total['emoji']}")
            st.success(f"**Eneatipo Global:** {enea_total} - {diag_total['nivel']} {diag_total['simbolo']}")
            st.write(f"**Conclusión General:** {diag_total['analisis']}")

            # Gráfico Visual
            fig, ax = plt.subplots(figsize=(10, 4))
            colores = ['#FF4B4B' if e <= 3 else '#00CC96' if e >= 7 else '#FFAA00' for e in df['Eneatipo']]
            ax.bar(df['Área'], df['Eneatipo'], color=colores, edgecolor='black')
            ax.set_ylim(0, 10)
            plt.xticks(rotation=15)
            st.pyplot(fig)

            # Mostrar tabla descriptiva en pantalla
            for r in res_data:
                st.markdown(f"**{r['Emoji']} {r['Área']} (Eneatipo {r['Eneatipo']}):** {r['Estado']}. *{r['Analisis']}*")

            # --- GENERACIÓN DEL DOCUMENTO WORD ---
            doc = Document()
            sec = doc.sections[0]
            sec.top_margin = sec.bottom_margin = Inches(0.4)
            sec.left_margin = sec.right_margin = Inches(0.5)

            # Título
            h = doc.add_heading('INFORME PSICOLÓGICO: HABILIDADES SOCIALES', 0)
            h.alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_paragraph(f"Paciente: {nombre}  |  Edad: {edad} años  |  Eneatipo Global: {enea_total}").bold = True

            # 1. Copia Fiel del Protocolo (2 Columnas)
            doc.add_heading('1. Protocolo de Evaluación (Autoinforme)', level=1)
            table = doc.add_table(rows=25, cols=4)
            table.style = 'Table Grid'
            for i in range(1, 26):
                table.cell(i-1, 0).text = f"{i}. {PREGUNTAS[i-1][:35]}..."
                table.cell(i-1, 1).text = str(respuestas[i])
                table.cell(i-1, 2).text = f"{i+25}. {PREGUNTAS[i+24][:35]}..."
                table.cell(i-1, 3).text = str(respuestas[i+25])
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs: p.runs[0].font.size = Pt(8)

            doc.add_page_break()

            # 2. Análisis Clínico
            doc.add_heading(f'2. Diagnóstico Global {diag_total["emoji"]}', level=1)
            doc.add_paragraph(f"Nivel Alcanzado: {diag_total['nivel']} (Eneatipo {enea_total}).")
            doc.add_paragraph(diag_total['analisis'])

            # Gráfico
            img_b = BytesIO()
            plt.savefig(img_b, format='png', bbox_inches='tight')
            doc.add_picture(img_b, width=Inches(6))

            # 3. Interpretación Extensa por Áreas
            doc.add_heading('3. Interpretación Extensa por Áreas', level=1)
            for r in res_data:
                p_area = doc.add_paragraph()
                p_area.add_run(f"■ {r['Área']} - Eneatipo {r['Eneatipo']} ({r['Estado']}): ").bold = True
                p_area.add_run(f"Evalúa {r['Desc'].lower()} ")
                p_area.add_run(f"\nAnálisis Clínico: {r['Analisis']}")

            # 4. Plan Terapéutico Detallado
            deficitarias = [r['Área'] for r in res_data if r['Eneatipo'] <= 3]
            doc.add_heading('4. Plan Terapéutico Estructurado', level=1)
            
            if deficitarias:
                doc.add_paragraph(f"Prioridad Clínica: Abordaje urgente en las áreas de {', '.join(deficitarias)}.")
                doc.add_paragraph("Fase 1: Psicoeducación y Reestructuración Cognitiva.\n  - Identificar pensamientos automáticos negativos ('Nadie me hará caso', 'Haré el ridículo').\n  - Modificar creencias limitantes y reducir la ansiedad anticipatoria.", style='List Bullet')
                doc.add_paragraph("Fase 2: Entrenamiento en Habilidades Sociales (EHS).\n  - Modelado: El terapeuta ejecuta la conducta deseada.\n  - Ensayo Conductual (Role-playing): El paciente practica la conducta en sesión simulada.", style='List Bullet')
                doc.add_paragraph("Fase 3: Feedback y Moldeamiento.\n  - Brindar retroalimentación positiva y corregir el lenguaje no verbal (contacto visual, postura).", style='List Bullet')
                doc.add_paragraph("Fase 4: Generalización.\n  - Asignar tareas de exposición gradual in vivo para practicar las habilidades en su entorno natural (escuela, familia, trabajo).", style='List Bullet')
            else:
                doc.add_paragraph("El paciente no presenta déficits clínicos significativos. El plan terapéutico debe enfocarse en el refuerzo positivo y el mantenimiento de las habilidades prosociales adquiridas, promoviendo su liderazgo en dinámicas grupales.")

            # Botón de Descarga
            w_buf = BytesIO()
            doc.save(w_buf)
            st.download_button("📥 DESCARGAR INFORME EXTENSO (WORD)", w_buf.getvalue(), f"Informe_Clinico_{nombre}.docx")
else:
    st.info("👈 Por favor, lea y acepte el consentimiento en la barra lateral izquierda para comenzar la evaluación.")
